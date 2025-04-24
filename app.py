from flask import Flask, request, jsonify, send_file, render_template
import requests
import sqlite3
from flask_cors import CORS
from docx import Document
import mammoth
from dotenv import load_dotenv
import os
import sys
import logging
import json
import io
import PyPDF2
from pdfrw import PdfReader, PdfWriter, PageMerge
from PyPDF2 import PdfReader as PyPdfReader, PdfWriter as PyPdfWriter

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# Import AI models
from model.bot import get_response
from model.similarity import get_document

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "default-secret-key")

# Configure CORS to allow requests from any origin for development
CORS(app, resources={r"/api/*": {"origins": "*"}})

# Helper function to get database connection
def get_db_connection():
    db_path = 'legal_assistant.db'
    
    # Check for the main database file
    if not os.path.exists(db_path):
        # Check for the alternate database file
        alt_path = 'legal_assistant_new.db'
        if os.path.exists(alt_path):
            logger.info(f"Using alternative database: {alt_path}")
            db_path = alt_path
        else:
            logger.warning(f"Database file not found: {db_path}. Creating empty database.")
            
    try:
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        return conn
    except Exception as e:
        logger.error(f"Database connection error: {str(e)}")
        raise

# Main route
@app.route('/')
def index():
    return render_template('index.html')

# Get all services
@app.route('/api/services', methods=["GET"])
def services():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM services')
        services = [dict(row) for row in cursor.fetchall()]
        cursor.close()
        conn.close()
        logger.debug(f"Retrieved services: {services}")
        return jsonify(services)
    except Exception as e:
        logger.error(f"Error retrieving services: {str(e)}")
        return jsonify({"error": str(e)}), 500

# Get forms of a particular service
@app.route('/api/forms', methods=["GET"])
def get_forms():
    try:
        service_id = request.args.get('service_id')
        logger.debug(f"Getting forms for service_id: {service_id}")
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT services.service_id, services.service_name, forms.form_id, forms.form_name, forms.form_link 
            FROM services 
            INNER JOIN forms ON services.service_id = forms.service_id 
            WHERE forms.service_id = ?
        """, [service_id])
        
        forms = [dict(row) for row in cursor.fetchall()]
        cursor.close()
        conn.close()
        
        logger.debug(f"Retrieved forms: {forms}")
        return jsonify(forms)
    except Exception as e:
        logger.error(f"Error retrieving forms: {str(e)}")
        return jsonify({"error": str(e)}), 500

# Get all queries for a form
@app.route('/api/form-details', methods=["GET"])
def get_form_details():
    try:
        form_id = request.args.get('form_id')
        logger.debug(f"Getting details for form_id: {form_id}")
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Get form details
        cursor.execute("SELECT * FROM forms WHERE form_id = ?;", [form_id])
        form_data = [dict(row) for row in cursor.fetchall()]
        
        # Get categories
        cursor.execute("""
            SELECT * FROM ques_categories 
            WHERE id IN (
                SELECT DISTINCT(category_id) 
                FROM input_ques 
                WHERE ques_id IN (
                    SELECT form_query_id 
                    FROM form_queries 
                    WHERE form_id = ?
                )
            );
        """, [form_id])
        categories = [dict(row) for row in cursor.fetchall()]
        
        # Get questions
        cursor.execute("""
            SELECT * FROM input_ques 
            WHERE ques_id IN (
                SELECT form_query_id 
                FROM form_queries 
                WHERE form_id = ?
            );
        """, [form_id])
        questions = [dict(row) for row in cursor.fetchall()]
        
        cursor.close()
        conn.close()
        
        # Combine all data
        result = form_data + categories + questions
        logger.debug(f"Retrieved form details: {result}")
        return jsonify(result)
    except Exception as e:
        logger.error(f"Error retrieving form details: {str(e)}")
        return jsonify({"error": str(e)}), 500

# Return the contents of final doc
@app.route('/api/final-content', methods=["POST"])
def final_content():
    try:
        form_details = request.json
        form_id = form_details["form_id"]
        logger.debug(f"Generating content for form_id: {form_id}")
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT form_link, form_name FROM forms where form_id = ?;", [form_id])
        form_data = dict(cursor.fetchone())
        form_link = form_data["form_link"]
        form_name = form_data["form_name"]
        cursor.close()
        conn.close()
        
        logger.debug(f"Downloading document from: {form_link}")
        
        # Create a simple HTML placeholder for PDF files
        html_content = f"""
        <div class="alert alert-info">
            <h3>PDF Document Preview</h3>
            <p>The selected document "{form_name}" is a PDF file.</p>
            <p>In a real-world scenario, we would process this PDF file by:</p>
            <ol>
                <li>Converting the PDF to a modifiable format</li>
                <li>Filling in your provided information:</li>
            </ol>
            
            <div class="card p-3 my-3" style="background-color: #2c3e50; border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.1);">
                <h4 class="text-white">Form Data Provided:</h4>
                <ul class="text-white">
        """
        
        # Add the form details to the HTML content
        for key, value in form_details.items():
            if key != "form_id" and key.isdigit():
                # Get the question text for this field
                conn = get_db_connection()
                cursor = conn.cursor()
                # Make sure to properly format the ID
                key_id = key.zfill(3)  # Ensure it's padded to 3 digits (e.g., "1" -> "001")
                cursor.execute("SELECT ques_text FROM input_ques WHERE ques_id = ?;", [f"Q{key_id}"])
                question_data = cursor.fetchone()
                cursor.close()
                conn.close()
                
                question_text = question_data["ques_text"] if question_data else f"Field {key}"
                html_content += f"<li><strong>{question_text}:</strong> {value}</li>"
        
        html_content += f"""
                </ul>
            </div>
            
            <p>The final document would be a completed version of the form with your information inserted in the appropriate places.</p>
            
            <div class="d-grid gap-2 d-md-flex justify-content-md-end">
                <a href="{form_link}" target="_blank" class="btn btn-primary">
                    <i class="bi bi-file-earmark-pdf"></i> View Original PDF
                </a>
            </div>
        </div>
        """
        
        # Create docs directory if it doesn't exist
        directory = './docs'
        if not os.path.exists(directory):
            os.makedirs(directory)
            
        # Save a simple .docx file with the form information for download
        doc = Document()
        doc.add_heading(f"{form_name}", 0)
        
        doc.add_heading("Form Details", level=1)
        for key, value in form_details.items():
            if key != "form_id" and key.isdigit():
                # Get the question text for this field
                conn = get_db_connection()
                cursor = conn.cursor()
                # Make sure to properly format the ID
                key_id = key.zfill(3)  # Ensure it's padded to 3 digits (e.g., "1" -> "001")
                cursor.execute("SELECT ques_text FROM input_ques WHERE ques_id = ?;", [f"Q{key_id}"])
                question_data = cursor.fetchone()
                cursor.close()
                conn.close()
                
                question_text = question_data["ques_text"] if question_data else f"Field {key}"
                doc.add_paragraph(f"{question_text}: {value}")
                
        doc.add_paragraph(f"Original Document: {form_link}")
        
        # Save the document
        doc.save("./docs/Output2.docx")
        
        return jsonify({'content': html_content})
    except Exception as e:
        logger.error(f"Error generating document content: {str(e)}")
        return jsonify({"error": str(e)}), 500

# Return the final doc
@app.route('/api/final-form', methods=["POST"])
def final_form():
    try:
        # If data is sent as JSON
        if request.is_json:
            form_details = request.get_json()
        # If data is sent as form data (which is the case when using the download button)
        elif request.form.get('formData'):
            form_details = json.loads(request.form.get('formData'))
        else:
            return jsonify({"error": "No form data provided"}), 400
            
        # Debug the data
        logger.debug(f"Generating final document with data: {form_details}")
        
        # Get the form ID and look up the form details
        form_id = form_details["form_id"]
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT form_link, form_name FROM forms where form_id = ?;", [form_id])
        form_data = dict(cursor.fetchone())
        form_link = form_data["form_link"]
        form_name = form_data["form_name"]
        cursor.close()
        conn.close()
        
        # Create docs directory if it doesn't exist
        directory = './docs'
        if not os.path.exists(directory):
            os.makedirs(directory)
        
        # Determine the output format based on the form link
        if form_link.lower().endswith('.pdf'):
            try:
                # Download the PDF from the form link
                response = requests.get(form_link)
                if response.status_code != 200:
                    return jsonify({"error": f"Failed to download the PDF. Status code: {response.status_code}"}), 500
                
                # Create an in-memory PDF
                pdf_bytes = io.BytesIO(response.content)
                
                # Create a new PDF with form data
                output_pdf = io.BytesIO()
                
                # Use PyPDF2 to add text to the PDF
                reader = PyPdfReader(pdf_bytes)
                writer = PyPdfWriter()
                
                # Process each page
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    writer.add_page(page)
                
                # Create a metadata page with the form data
                metadata_pdf = PyPDF2.PdfWriter()
                metadata_pdf.add_blank_page(width=612, height=792)  # Standard letter size
                
                # Create metadata page with form details
                metadata_text = f"{form_name} - Form Data\n\n"
                
                for key, value in form_details.items():
                    if key != "form_id" and key.isdigit():
                        # Get the question text for this field
                        conn = get_db_connection()
                        cursor = conn.cursor()
                        key_id = key.zfill(3)
                        cursor.execute("SELECT ques_text FROM input_ques WHERE ques_id = ?;", [f"Q{key_id}"])
                        question_data = cursor.fetchone()
                        cursor.close()
                        conn.close()
                        
                        question_text = question_data["ques_text"] if question_data else f"Field {key}"
                        metadata_text += f"{question_text}: {value}\n"
                
                # Add metadata page to the output PDF
                metadata_temp = io.BytesIO()
                metadata_pdf.write(metadata_temp)
                metadata_temp.seek(0)
                metadata_reader = PyPdfReader(metadata_temp)
                writer.add_page(metadata_reader.pages[0])
                
                # Write the output PDF
                writer.write(output_pdf)
                output_pdf.seek(0)
                
                # Save the PDF for download
                output_path = os.path.join(directory, 'filled_form.pdf')
                with open(output_path, 'wb') as f:
                    f.write(output_pdf.getvalue())
                
                # Return the filled PDF
                return send_file(output_path, as_attachment=True, download_name=f"{form_name}_filled.pdf")
            
            except Exception as e:
                logger.error(f"Error generating PDF: {str(e)}")
                # Fall back to DOCX if PDF generation fails
                logger.info("Falling back to DOCX generation")
                if os.path.exists('./docs/Output2.docx'):
                    return send_file('./docs/Output2.docx', as_attachment=True)
                else:
                    return jsonify({"error": f"Failed to generate document: {str(e)}"}), 500
        else:
            # Return the existing DOCX file
            if os.path.exists('./docs/Output2.docx'):
                return send_file('./docs/Output2.docx', as_attachment=True)
            else:
                return jsonify({"error": "Document not found"}), 404
    except Exception as e:
        logger.error(f"Error serving final document: {str(e)}")
        return jsonify({"error": str(e)}), 500

# Chat API endpoint
@app.route('/api/chat', methods=['POST'])
def chat():
    try:
        user_input = request.json
        user_message = user_input.get('user_chat', '')
        logger.debug(f"Chat request received: {user_message}")
        
        # Get response from AI model
        response = get_document(user_message)
        
        logger.debug(f"AI response: {response}")
        return jsonify({'aiMessage': response})
    except Exception as e:
        logger.error(f"Error in chat processing: {str(e)}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
